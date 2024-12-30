from typing import List, Dict
from docx import Document
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import openai

class DocumentGenerator:
    @staticmethod
    def generate_summary(messages: List[Dict[str, str]], goal: str) -> str:
        """Generate a summary document from the conversation.
        
        Args:
            messages: List of conversation messages
            goal: The original goal/question
            
        Returns:
            str: Generated summary text
        """
        summary_prompt = {
            "role": "system",
            "content": f"""Summarize the entire conversation, focusing on the following:

        * **Executive Summary:** 
            - Briefly describe the main topic and purpose of the conversation. 
            - Summarize the key outcomes and decisions reached.

        * **Key Findings & Insights:** 
            - Highlight the most important information, conclusions, and agreements made during the discussion. 
            - Include any valuable insights or new perspectives gained.

        * **Action Items:** 
            - List any specific actions, decisions, or next steps that need to be taken. 
            - Include who is responsible for each action and any associated deadlines.

        Keep the summary concise, objective, and easy to understand. Avoid generic statements and focus on the most relevant and actionable information.

        **Goal of Conversation:**
        {goal}

        **Conversation:**
        """
        }

        
        response = openai.chat.completions.create(
            model="gpt-4",
            messages=[summary_prompt] + messages,
            temperature=0.6,
            max_tokens=1000,
            stream=False
        )
        return response.choices[0].message.content


    @staticmethod
    def create_document(summary: str, goal: str, messages: List[Dict[str, str]], agent1_name: str, agent2_name: str) -> io.BytesIO:
        doc = Document()

        header = doc.add_heading(f"{agent1_name} X {agent2_name} Conversation", 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        goal_para = doc.add_paragraph()
        goal_run = goal_para.add_run('Innovation Goal: ')
        goal_run.bold = True
        goal_para.add_run(goal)

        doc.add_paragraph()

        lines = summary.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('#'):
                heading_level = min(line.count('#'), 4)
                heading_text = line.lstrip('#').strip()
                doc.add_heading(heading_text, level=heading_level)
            elif line:
                para = doc.add_paragraph()
                format_text(para, line)

        doc.add_heading('Full Dialogue', level=1)
        for i, msg in enumerate(messages):
            p = doc.add_paragraph()
            if i % 2 == 0: 
                speaker_run = p.add_run(f"{agent1_name}: \n")
            else:
                speaker_run = p.add_run(f"{agent2_name}: \n")
            speaker_run.bold = True

            message_content = msg["content"]
            process_markdown_content(doc, message_content)

            if msg != messages[-1]:
                doc.add_paragraph('---')

        doc.add_paragraph('=' * 50)
        footer = doc.add_paragraph('Generated by GreenPill Network x Kevin Owocki Dialogue System')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io
    
def process_markdown_content(doc, content):
    """
    Processes markdown-like content in messages to handle headings and bold text.
    """
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if line.startswith('#'):
            heading_level = min(line.count('#'), 4)
            heading_text = line.lstrip('#').strip()
            doc.add_heading(heading_text, level=heading_level)
        elif line:
            para = doc.add_paragraph()
            format_text(para, line)

def format_text(paragraph, content):
    """
    Formats text to handle **<TEXT>** for bold formatting and removes the ** markers.
    """
    parts = re.split(r'(\*\*.*?\*\*)', content)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)
