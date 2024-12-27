from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import streamlit as st
import openai

def generate_summary_doc(messages, goal):
    """Generate a summary document from the conversation.
    
    Args:
        messages: List of conversation messages
        goal: The original goal/question
        
    Returns:
        str: Generated summary text
    """
    summary_prompt = {
        "role": "system",
        "content": """Create a concise but comprehensive Web3 x Regenerative Future summary. Keep each section focused and specific:
        
        # Executive Summary
        [One paragraph overview of key innovations]
        
        # Technical Innovations
        - Core mechanisms proposed
        - Novel combinations discovered
        - Technical challenges addressed
        
        # Implementation Framework
        - Key milestones
        - Technical requirements
        - Resource needs
        
        # Impact Metrics
        - Environmental KPIs
        - Social impact measures
        - Economic sustainability indicators
        
        Keep all content specific and actionable. Avoid generic statements."""
    }
    
    try:
        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[summary_prompt] + messages,
            temperature=0.7,
            max_tokens=1000,
            stream=False
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error generating summary: {str(e)}")
        return None

def create_formatted_docx(summary, goal, messages):
    """
        This function creates a formatted docx document from the conversation messages.
        
        Args :
            summary (str) : The generated summary.
            goal (str) : The innovation goal.
            messages (List[Dict[str, str]]) : The conversation messages.
            
        Returns :
            io.BytesIO : The generated document.
    """
    try:
        doc = Document()
        
        header = doc.add_heading('Web3 x Regenerative Future Dialogue', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        goal_para = doc.add_paragraph()
        goal_run = goal_para.add_run('Innovation Goal: ')
        goal_run.bold = True
        goal_para.add_run(goal)
        
        doc.add_paragraph()
        
        sections = summary.split('#')
        for section in sections[1:]:
            lines = section.strip().split('\n')
            if lines:
                doc.add_heading(lines[0].strip(), level=1)
                content = '\n'.join(lines[1:]).strip()
                if content:
                    para = doc.add_paragraph()
                    para.add_run(content)
        
        doc.add_heading('Full Dialogue', level=1)
        for msg in messages:
            p = doc.add_paragraph()
            if msg["role"] == "user":
                p.add_run("🌱 GreenPillAI: ").bold = True
            else:
                p.add_run("👨‍💻 Kevin: ").bold = True
            p.add_run(msg["content"])
            doc.add_paragraph('---')
        
        doc.add_paragraph('=' * 50)
        footer = doc.add_paragraph('Generated by GreenPill Network x Kevin Owocki Dialogue System')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io
        
    except Exception as e:
        st.error(f"Error creating document: {str(e)}")
        return None